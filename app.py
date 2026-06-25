import io
import os
import time

import docx
import pandas as pd
import plotly.express as px
import streamlit as st
from fpdf import FPDF

import analyzer
import database
import document_parser
import html_exporter

# --- Password Gate ---
# Simple shared-password protection for internal team deployment. Not a
# substitute for real auth/SSO, but enough to keep an unlisted link from
# being usable by anyone who stumbles onto the URL.
APP_PASSWORD = os.environ.get("APP_PASSWORD", "")


def _check_password() -> bool:
    if not APP_PASSWORD:
        # No password configured (e.g. local dev) - skip the gate entirely.
        return True
    if st.session_state.get("password_ok"):
        return True

    st.title("🛡️ TPRM Report Analyzer")
    pw = st.text_input("Team password", type="password")
    if st.button("Enter"):
        if pw == APP_PASSWORD:
            st.session_state["password_ok"] = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    return False


st.set_page_config(page_title="TPRM Report Analyzer", layout="wide", page_icon="🛡️")

if not _check_password():
    st.stop()

# Initialize local feedback DB
database.init_db()

# --- Custom Styling ---
st.markdown("""
    <style>
    .main {background-color: #0e1117;}
    .stButton>button {width: 100%;}
    </style>
""", unsafe_allow_html=True)

# --- State Initialization ---
if 'analysis_done' not in st.session_state:
    st.session_state['analysis_done'] = False
if 'results' not in st.session_state:
    st.session_state['results'] = []
if 'errors' not in st.session_state:
    st.session_state['errors'] = []

st.title("🛡️ TPRM Report Analyzer")
st.markdown("**Comprehensive, interactive Third-Party Risk Management evaluation powered by AI.**")
st.write("Upload vendor risk files, provide custom instructions, and generate leadership-ready reports.")

# --- 1 & 2. FILE INPUT & CUSTOM PROMPT INTERFACE ---
with st.sidebar:
    st.header("🤖 AI Engine")
    if os.environ.get("ANTHROPIC_API_KEY"):
        st.success(f"Connected — using {analyzer.MODEL_NAME}")
    else:
        st.error("ANTHROPIC_API_KEY is not configured. Analysis will fail until it's set.")

    st.header("1. Upload Vendor Files")
    uploaded_files = st.file_uploader(
        "Upload Reports (.pdf, .docx, .xlsx)",
        type=['pdf', 'docx', 'xlsx'],
        accept_multiple_files=True
    )

    st.header("2. Custom Prompt Section")
    custom_prompt = st.text_area(
        "Enter custom analysis instructions:",
        placeholder="e.g., 'Focus heavily on GDPR compliance,' or 'Filter out any findings older than 2024'",
        help="Claude adapts its analysis based on these instructions."
    )

    analyze_btn = st.button("🚀 Run Analysis", type="primary")


# --- CATEGORY SCORING (derived from real findings, not random) ---
RISK_CATEGORIES = {
    "Access Control": ["access", "authentication", "mfa", "password", "privilege", "credential", "login", "identity"],
    "Data Privacy": ["data privacy", "pii", "gdpr", "encryption", "data protection", "personal data", "confidential"],
    "Patch Management": ["patch", "vulnerability", "outdated", "update", "cve", "unpatched", "version"],
    "Compliance": ["compliance", "audit", "soc 2", "iso", "policy", "regulation", "certification", "control"],
}

RISK_BASE_SCORE = {"Critical": 85, "High": 65, "Medium": 45, "Low": 20}


def compute_category_scores(result: dict) -> dict:
    """
    Derives a 0-100 score per category from the AI's actual findings text,
    rather than random noise. Categories with keyword hits in the findings
    score higher (capped at 95); others fall back to a baseline derived
    from the overall risk level, so the chart still looks reasonable for
    vendors where findings don't map cleanly onto these four buckets.
    """
    base = RISK_BASE_SCORE.get(result.get("risk_score", "Medium"), 45)
    findings_text = " ".join(result.get("findings", [])).lower()

    scores = {}
    for category, keywords in RISK_CATEGORIES.items():
        hits = sum(findings_text.count(kw) for kw in keywords)
        score = min(95, base + hits * 12)
        scores[category] = score
    return scores


def _safe_result(res: dict) -> dict:
    """
    Defense-in-depth guard: analyzer._normalize_result already sanitizes
    AI output, but this guarantees that any result dict flowing through the
    rest of the app (export functions, charts, sorting) can never have a
    None/missing risk_score or vendor, even if it arrived via some other
    path. `.get(key, default)` only supplies the default when the key is
    MISSING, not when it's present-but-None - so we use `or` here instead.
    """
    res = dict(res)
    res["vendor"] = res.get("vendor") or "Unknown Vendor"
    res["risk_score"] = res.get("risk_score") or "Medium"
    res["summary"] = res.get("summary") or ""
    res["scope"] = res.get("scope") or ""
    res["postscript"] = res.get("postscript") or ""
    res["findings"] = res.get("findings") or []
    res["next_steps"] = res.get("next_steps") or []
    return res


# --- 3. ANALYSIS ENGINE ---
def run_live_analysis(files, prompt):
    results = []
    max_retries = 3

    for file in files:
        with st.spinner(f"Extracting text from {file.name}..."):
            doc_text, parse_error = document_parser.parse_file(file)

        if parse_error:
            st.session_state['errors'].append(f"Skipped {file.name}: {parse_error}")
            continue

        res = {"error": "Analysis did not run."}
        for attempt in range(max_retries):
            with st.spinner(f"Analyzing {file.name} with Claude ({analyzer.MODEL_NAME})... (attempt {attempt + 1}/{max_retries})"):
                res = analyzer.analyze_vendor_document(doc_text, prompt, max_retries=1)

            if "error" not in res:
                break

            is_transient = any(s in res["error"].lower() for s in ["rate limit", "connect", "empty response", "overload"])
            if is_transient and attempt < max_retries - 1:
                wait = min(5 * (attempt + 1), 30)
                st.warning(f"Transient issue analyzing {file.name} ({res['error']}). Retrying in {wait}s...")
                time.sleep(wait)
                continue
            else:
                break

        if "error" in res:
            st.session_state['errors'].append(f"Error analyzing {file.name}: {res['error']}")
            continue

        res["filename"] = file.name
        res = _safe_result(res)
        res["category_scores"] = compute_category_scores(res)
        results.append(res)

    risk_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    results.sort(key=lambda x: risk_order.get(x.get("risk_score", "Medium"), 2))
    return results


# --- 4. MULTI-FORMAT OUTPUT GENERATION ---
def generate_docx(results):
    doc = docx.Document()
    doc.add_heading('Executive Leadership Summary', 0)
    doc.add_paragraph(f"Total Vendors Assessed: {len(results)}")

    if not results:
        doc.add_paragraph("No vendor results were available for this report.")
    else:
        doc.add_heading('Vendor Breakdowns', 1)
        for raw_res in results:
            res = _safe_result(raw_res)
            doc.add_heading(f"{res.get('vendor', 'Unknown')} - {res.get('risk_score', 'Medium')} Risk", 2)
            doc.add_paragraph(res.get('summary', ''))
            if res.get('findings'):
                doc.add_paragraph("Critical Findings:")
                for finding in res['findings']:
                    doc.add_paragraph(f"{finding}", style='List Bullet')
            if res.get('next_steps'):
                doc.add_paragraph("Next Steps:")
                for step in res['next_steps']:
                    doc.add_paragraph(f"{step}", style='List Bullet')
            if res.get('postscript'):
                doc.add_paragraph(f"Leadership Postscript: {res.get('postscript', '')}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _pdf_clean_text(text):
    """fpdf's core fonts only support latin-1. Map common 'smart' punctuation
    to ASCII equivalents first, then drop anything else unsupported rather
    than crashing the export or silently corrupting vendor names."""
    if not text:
        return ""
    text = str(text)
    replacements = {
        '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u00A0': ' ',
        '\u2022': '-', '\u2192': '->',
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode('latin-1', 'ignore').decode('latin-1')


def generate_pdf(results):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    bg_dark = (14, 17, 23)
    accent_blue = (79, 139, 249)
    text_main = (30, 30, 30)
    text_muted = (100, 100, 100)

    # 1. Header Banner
    pdf.set_fill_color(*bg_dark)
    pdf.rect(0, 0, 210, 40, 'F')

    pdf.set_xy(10, 12)
    pdf.set_font("Arial", size=24, style="B")
    pdf.set_text_color(255, 255, 255)
    pdf.cell(190, 10, txt="TPRM Executive Analysis", ln=1, align='C')

    pdf.set_font("Arial", size=11, style="I")
    pdf.set_text_color(180, 180, 180)
    pdf.cell(190, 8, txt="Highly Confidential - Automated AI Assessment", ln=1, align='C')

    pdf.set_xy(10, 50)

    # 2. Portfolio Overview
    pdf.set_font("Arial", size=16, style="B")
    pdf.set_text_color(*accent_blue)
    pdf.cell(190, 10, txt="Portfolio Overview", ln=1, border='B')
    pdf.ln(5)

    pdf.set_font("Arial", size=12, style="B")
    pdf.set_text_color(*text_main)
    pdf.cell(95, 8, txt=f"Total Vendors Assessed: {len(results)}", ln=0)

    critical_count = sum(1 for r in results if (r.get('risk_score') or 'Medium') == 'Critical')
    if critical_count > 0:
        pdf.set_text_color(255, 75, 75)
    else:
        pdf.set_text_color(*text_main)
    pdf.cell(95, 8, txt=f"Critical Vulnerabilities: {critical_count}", ln=1)
    pdf.ln(10)

    if not results:
        pdf.set_font("Arial", size=12, style="I")
        pdf.set_text_color(*text_muted)
        pdf.multi_cell(190, 8, txt="No vendor results were available for this report.")
        return bytes(pdf.output())

    # 3. Vendor Cards
    for raw_res in results:
        res = _safe_result(raw_res)
        # Reserve roughly the space a card needs; if it won't fit on the
        # current page, start a fresh page rather than letting fpdf's
        # auto-page-break split a card awkwardly mid-section.
        estimated_height = 70 + 6 * (len(res.get('findings', [])) + len(res.get('next_steps', [])))
        if pdf.get_y() + min(estimated_height, 240) > 270:
            pdf.add_page()

        vendor = _pdf_clean_text(res.get('vendor', 'Unknown'))
        risk = res.get('risk_score', 'Medium')

        pdf.set_fill_color(240, 242, 246)
        pdf.cell(190, 12, txt="", ln=0, fill=True)
        pdf.set_x(10)

        pdf.set_font("Arial", size=14, style="B")
        pdf.set_text_color(*text_main)
        pdf.cell(140, 12, txt=f" {vendor}", ln=0)

        pdf.set_font("Arial", size=11, style="B")
        if risk == "Critical":
            pdf.set_text_color(255, 255, 255)
            pdf.set_fill_color(255, 75, 75)
        elif risk == "High":
            pdf.set_text_color(255, 255, 255)
            pdf.set_fill_color(255, 166, 0)
        elif risk == "Medium":
            pdf.set_text_color(30, 30, 30)
            pdf.set_fill_color(255, 222, 36)
        else:
            pdf.set_text_color(255, 255, 255)
            pdf.set_fill_color(0, 209, 78)

        pdf.cell(45, 8, txt=f"{risk.upper()} RISK", ln=1, align='C', fill=True)
        pdf.set_x(10)
        pdf.ln(4)

        pdf.set_font("Arial", size=11, style="B")
        pdf.set_text_color(*accent_blue)
        pdf.set_x(10)
        pdf.cell(190, 8, txt="Context & Scope", ln=1)

        pdf.set_font("Arial", size=10)
        pdf.set_text_color(*text_muted)
        pdf.set_x(10)
        pdf.multi_cell(190, 6, txt=_pdf_clean_text(res.get('summary', '')))
        pdf.set_x(10)
        pdf.multi_cell(190, 6, txt=_pdf_clean_text(res.get('scope', '')))
        pdf.ln(4)

        pdf.set_font("Arial", size=11, style="B")
        pdf.set_text_color(*accent_blue)
        pdf.set_x(10)
        pdf.cell(190, 8, txt="Critical Findings", ln=1)

        pdf.set_font("Arial", size=10)
        pdf.set_text_color(*text_main)
        for finding in res.get('findings', []):
            pdf.set_x(10)
            pdf.multi_cell(190, 6, txt=_pdf_clean_text(f"- {finding}"))
        pdf.ln(4)

        pdf.set_font("Arial", size=11, style="B")
        pdf.set_text_color(*accent_blue)
        pdf.set_x(10)
        pdf.cell(190, 8, txt="Remediation Next Steps", ln=1)

        pdf.set_font("Arial", size=10)
        pdf.set_text_color(*text_main)
        for step in res.get('next_steps', []):
            pdf.set_x(10)
            pdf.multi_cell(190, 6, txt=_pdf_clean_text(f"- {step}"))
        pdf.ln(4)

        pdf.set_fill_color(235, 242, 255)
        pdf.set_text_color(10, 40, 100)
        pdf.set_font("Arial", size=10, style="B")
        pdf.set_x(10)
        pdf.cell(190, 8, txt=" TPRM LEADERSHIP POSTSCRIPT", ln=1, fill=True)
        pdf.set_font("Arial", size=10)
        pdf.set_x(10)
        pdf.multi_cell(190, 6, txt=_pdf_clean_text(res.get('postscript', '')), fill=True)

        pdf.ln(10)

    return bytes(pdf.output())


# Main Execution Flow
if analyze_btn:
    if not uploaded_files:
        st.error("Please upload at least one vendor file.")
    elif not os.environ.get("ANTHROPIC_API_KEY"):
        st.error("ANTHROPIC_API_KEY is not configured. Set it in your environment before running analysis.")
    else:
        st.session_state['errors'] = []
        st.session_state['results'] = run_live_analysis(uploaded_files, custom_prompt)
        st.session_state['analysis_done'] = True

if st.session_state.get('analysis_done'):
    if st.session_state.get('errors'):
        for err in st.session_state['errors']:
            st.error(err)

    results = st.session_state.get('results', [])
    if not results:
        st.warning("No vendor results to display. See errors above, or try re-running the analysis.")
        st.stop()

    st.success("Analysis Complete!")

    # --- EXECUTIVE LEADERSHIP SUMMARY ---
    st.header("Executive Leadership Summary")
    col1, col2 = st.columns([1, 2])
    with col1:
        st.metric("Total Vendors Assessed", len(results))
        risk_counts = pd.DataFrame([r.get('risk_score', 'Medium') for r in results], columns=['Risk']).value_counts().reset_index()
        risk_counts.columns = ['Risk', 'Count']
        st.dataframe(risk_counts, hide_index=True, use_container_width=True)

    with col2:
        if not risk_counts.empty:
            fig = px.pie(risk_counts, values='Count', names='Risk', title="Master Risk Distribution",
                         color='Risk', color_discrete_map={"Critical": "red", "High": "orange", "Medium": "yellow", "Low": "green"})
            st.plotly_chart(fig, use_container_width=True)

    st.markdown("---")

    # --- DETAILED VENDOR BREAKDOWN ---
    st.header("Detailed Vendor-by-Vendor Breakdown")
    for i, res in enumerate(results):
        vendor_name = res.get('vendor', 'Unknown')
        with st.expander(f"{res.get('risk_score', 'Medium')} RISK: {vendor_name}", expanded=True):
            st.write(f"**Context:** {res.get('summary', '')}")
            st.write(f"**Scope:** {res.get('scope', '')}")

            col_f, col_c = st.columns(2)
            with col_f:
                st.markdown("**Critical Findings:**")
                for f in res.get('findings', []):
                    st.markdown(f"- {f}")

                st.markdown("**Next Steps & Remediation:**")
                for s in res.get('next_steps', []):
                    st.markdown(f"- {s}")

            with col_c:
                # Chart driven by real keyword analysis of the findings text,
                # not random data.
                category_scores = res.get("category_scores") or compute_category_scores(res)
                df_chart = pd.DataFrame({
                    "Category": list(category_scores.keys()),
                    "Vulnerability Score": list(category_scores.values()),
                })
                v_fig = px.bar(df_chart, x='Category', y='Vulnerability Score', title="Vulnerability Breakdown",
                                range_y=[0, 100])
                st.plotly_chart(v_fig, use_container_width=True, key=f"chart_{vendor_name}_{i}")

            st.info(f"**Leadership Postscript:** {res.get('postscript', '')}")

            # --- 5. INTERACTIVE USER FEEDBACK ---
            st.markdown("---")
            st.write("**Was this analysis accurate?**")
            fcol1, fcol2, _ = st.columns([1, 1, 8])

            with fcol1:
                if st.button("👍 Thumbs Up", key=f"up_{vendor_name}_{i}"):
                    ok = database.log_feedback(vendor_name, res.get("filename", ""), True)
                    if ok:
                        st.toast(f"Feedback logged for {vendor_name}.")
                    else:
                        st.toast("Could not save feedback (local storage issue).", icon="⚠️")
            with fcol2:
                if st.button("👎 Thumbs Down", key=f"down_{vendor_name}_{i}"):
                    st.session_state[f"dialog_{vendor_name}_{i}"] = True

            if st.session_state.get(f"dialog_{vendor_name}_{i}", False):
                with st.form(key=f"form_{vendor_name}_{i}"):
                    st.write("Help improve future analyses:")
                    feedback_text = st.text_area("Explicit comments (optional):", placeholder="e.g., 'The risk rating for Vendor X was too high'")
                    submitted = st.form_submit_button("Submit")
                    if submitted:
                        ok = database.log_feedback(vendor_name, res.get("filename", ""), False, feedback_text)
                        if ok:
                            st.toast("Feedback logged.")
                        else:
                            st.toast("Could not save feedback (local storage issue).", icon="⚠️")
                        st.session_state[f"dialog_{vendor_name}_{i}"] = False
                        st.rerun()

    # --- DOWNLOAD OPTIONS ---
    st.markdown("---")
    st.header("Export Final Report")

    try:
        pdf_data = generate_pdf(results)
        st.download_button(
            label="📄 Download as Executive PDF Report (.pdf)",
            data=pdf_data,
            file_name="TPRM_Executive_Report.pdf",
            mime="application/pdf",
            type="primary",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Could not generate PDF report: {e}")

    st.markdown("<br><p style='text-align: center; color: gray;'>Alternative Formats</p>", unsafe_allow_html=True)

    col_d1, col_d2 = st.columns(2)
    with col_d1:
        try:
            html_data = html_exporter.generate_html_report(results)
            st.download_button(
                label="🌐 Download as Interactive HTML (.html)",
                data=html_data,
                file_name="TPRM_Interactive_Report.html",
                mime="text/html",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Could not generate HTML report: {e}")
    with col_d2:
        try:
            docx_data = generate_docx(results)
            st.download_button(
                label="📄 Download as Word (.docx)",
                data=docx_data,
                file_name="TPRM_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Could not generate Word report: {e}")
