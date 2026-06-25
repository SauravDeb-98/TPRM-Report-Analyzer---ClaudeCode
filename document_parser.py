"""
Document parsing for uploaded vendor risk files (.pdf, .docx, .xlsx).

Returns a (text, error) tuple instead of burying parse failures inside the
extracted text, so callers can decide whether to skip analysis entirely
rather than asking the AI to "analyze" an error message.
"""

import io

import docx
import fitz  # PyMuPDF
import pandas as pd

# Guardrails against pathological inputs (corrupted files, accidental huge
# uploads, zip-bomb-style xlsx, etc.)
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25 MB
MAX_PDF_PAGES = 500
MAX_EXCEL_ROWS_PER_SHEET = 5000


def parse_file(uploaded_file):
    """
    Parses an uploaded file based on its extension.

    Args:
        uploaded_file: a Streamlit UploadedFile (or any object exposing
            .name, .read(), and .seek()).

    Returns:
        (text_content: str, error: str | None)
        If error is not None, text_content may still contain partial
        content but should generally not be sent to the analyzer.
    """
    name = getattr(uploaded_file, "name", "uploaded_file")
    file_extension = name.split(".")[-1].lower() if "." in name else ""

    try:
        file_bytes = uploaded_file.read()
    except Exception as e:
        return "", f"Could not read file '{name}': {e}"
    finally:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass

    if not file_bytes:
        return "", f"'{name}' is empty (0 bytes)."

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        size_mb = len(file_bytes) / (1024 * 1024)
        return "", f"'{name}' is {size_mb:.1f} MB, which exceeds the {MAX_FILE_SIZE_BYTES // (1024*1024)} MB limit."

    text_content = f"Filename: {name}\n"

    if file_extension == "pdf":
        return _parse_pdf(file_bytes, name, text_content)
    elif file_extension == "docx":
        return _parse_docx(file_bytes, name, text_content)
    elif file_extension == "xlsx":
        return _parse_xlsx(file_bytes, name, text_content)
    else:
        return "", f"Unsupported file format '.{file_extension}' for '{name}'. Supported: .pdf, .docx, .xlsx"


def _parse_pdf(file_bytes, name, text_content):
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        return "", f"'{name}' could not be opened as a PDF (corrupted or invalid file): {e}"

    try:
        if doc.is_encrypted:
            # Try an empty password (common for "owner-locked but readable" PDFs);
            # if that fails, surface a clear error rather than silently extracting nothing.
            if not doc.authenticate(""):
                doc.close()
                return "", f"'{name}' is password-protected and could not be opened automatically."

        if doc.page_count == 0:
            doc.close()
            return "", f"'{name}' contains no pages."

        if doc.page_count > MAX_PDF_PAGES:
            doc.close()
            return "", f"'{name}' has {doc.page_count} pages, exceeding the {MAX_PDF_PAGES}-page limit."

        extracted_any = False
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                extracted_any = True
            text_content += page_text + "\n"
        doc.close()

        if not extracted_any:
            return text_content, (
                f"'{name}' produced no extractable text (likely a scanned/image-only PDF "
                f"without OCR). Analysis would be based on an empty document."
            )
        return text_content, None

    except Exception as e:
        try:
            doc.close()
        except Exception:
            pass
        return "", f"Error while extracting text from '{name}': {e}"


def _parse_docx(file_bytes, name, text_content):
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        return "", f"'{name}' could not be opened as a Word document (corrupted or invalid file): {e}"

    try:
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                text_content += " | ".join(cell.text for cell in row.cells) + "\n"

        if len(text_content.strip()) <= len(f"Filename: {name}".strip()):
            return text_content, f"'{name}' appears to contain no readable text or tables."
        return text_content, None
    except Exception as e:
        return "", f"Error while extracting text from '{name}': {e}"


def _parse_xlsx(file_bytes, name, text_content):
    try:
        excel_data = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    except Exception as e:
        return "", f"'{name}' could not be opened as an Excel file (corrupted or invalid file): {e}"

    try:
        if not excel_data:
            return "", f"'{name}' contains no sheets."

        any_data = False
        for sheet_name, df in excel_data.items():
            text_content += f"--- Sheet: {sheet_name} ---\n"
            if df.empty:
                text_content += "(empty sheet)\n"
                continue
            any_data = True
            if len(df) > MAX_EXCEL_ROWS_PER_SHEET:
                text_content += (
                    f"(showing first {MAX_EXCEL_ROWS_PER_SHEET} of {len(df)} rows)\n"
                )
                df = df.head(MAX_EXCEL_ROWS_PER_SHEET)
            text_content += df.to_string(index=False) + "\n"

        if not any_data:
            return text_content, f"'{name}' has no data in any sheet."
        return text_content, None
    except Exception as e:
        return "", f"Error while extracting text from '{name}': {e}"
